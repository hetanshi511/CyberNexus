from typing import TypedDict, Dict, Any, List, Optional
from langgraph.graph import StateGraph, END
import json
import re
import asyncio
import requests
import logging
import os
import docx
from requests.auth import HTTPBasicAuth
from utils.llm import get_llm
from utils.db import (
    get_ticket_hash,
    load_report_from_db,
    upsert_ticket, upsert_report,
    compute_ticket_hash,
)
from utils.attachment import process_ticket_attachments

logger = logging.getLogger(__name__)


# -------------------- STATE --------------------

class ComplianceState(TypedDict):
    project_key: str
    ticket_id: str          # Optional single ticket
    jira_domain: str
    jira_email: str
    jira_token: str

    tickets: List[Dict[str, Any]]        # List of extracted ticket data
    project_report: List[Dict[str, Any]] # List of analysis results
    final_error: str
    tenant_id: str   # Firebase UID — multi-tenant


# -------------------- HELPERS --------------------

def extract_text_from_adf(adf):
    """Minimal extractor for Atlassian Document Format. Extracts plain text recursively."""
    if not adf:
        return ""
    if isinstance(adf, str):
        return adf

    text_parts = []

    def recurse(node):
        if isinstance(node, dict):
            if node.get("type") == "text":
                text_parts.append(node.get("text", ""))
            for value in node.values():
                recurse(value)
        elif isinstance(node, list):
            for item in node:
                recurse(item)

    recurse(adf)
    return " ".join(text_parts).strip()


def read_compliance_table():
    checklist_items = []
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        table_path = os.path.join(base_dir, "Table.docx")

        if not os.path.exists(table_path):
            logger.warning(f"Table.docx not found at {table_path}")
            return []

        doc = docx.Document(table_path)
        if not doc.tables:
            return []

        table = doc.tables[0]
        # Skip header (row 0)
        for row in table.rows[1:]:
            cells = row.cells
            # Expecting at least 2 columns: Category, Specific Check
            if len(cells) >= 2:
                cat = cells[0].text.strip()
                check = cells[1].text.strip()
                if cat or check:
                    checklist_items.append({"category": cat, "check": check})

        return checklist_items

    except Exception as e:
        logger.error(f"Failed to read Compliance Table: {e}")
        return []


def safe_json_parse(content: str) -> Optional[dict]:
    """
    Robust JSON parser.
    1. Try direct parse on raw content.
    2. Strip markdown code fences and retry.
    3. Regex-extract first {...} block and retry.
    Returns None if all attempts fail.
    """
    # Attempt 1: direct parse
    try:
        return json.loads(content)
    except (json.JSONDecodeError, TypeError):
        pass

    # Attempt 2: strip code fences
    cleaned = re.sub(r"```json|```", "", content).strip()
    try:
        return json.loads(cleaned)
    except (json.JSONDecodeError, TypeError):
        pass

    # Attempt 3: regex extract first {...} block
    match = re.search(r"\{.*\}", cleaned, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except (json.JSONDecodeError, TypeError):
            pass

    return None


def _error_analysis(reason: str) -> dict:
    """Safe fallback analysis when LLM fails entirely."""
    return {
        "alignment_status": "Error",
        "severity": "Low",
        "completion_percentage": 0,
        "compliance_gaps": [f"AI response parse failure: {reason}"],
        "recommended_actions": ["Retry analysis"],
        "compliance_checklist": []
    }


# -------------------- NODE 1: FETCH (PAGINATED via JQL) --------------------

def fetch_tickets(state: ComplianceState):
    logger.info(
        f"[Fetch] Starting — Project={state.get('project_key')}, "
        f"Ticket={state.get('ticket_id')}"
    )

    domain = state["jira_domain"]
    email = state["jira_email"]
    token = state["jira_token"]

    if not domain.startswith("http"):
        domain = f"https://{domain}"

    # ---- Resolve custom field IDs dynamically ----
    def get_field_id_by_name(name_query):
        try:
            res = requests.get(
                f"{domain}/rest/api/3/field",
                auth=HTTPBasicAuth(email, token),
                headers={"Accept": "application/json"},
                timeout=30
            )
            if res.status_code == 200:
                for f in res.json():
                    if name_query.lower() in f["name"].lower():
                        return f["id"]
        except Exception as e:
            logger.error(f"[Fetch] Error resolving field '{name_query}': {e}")
        return None

    approver_field = get_field_id_by_name("Approver")
    reviewer_field = get_field_id_by_name("Reviewer")
    logger.info(f"[Fetch] Resolved — Approver: {approver_field}, Reviewer: {reviewer_field}")

    # ---- Build fields param string for Agile API ----
    fields_list = [
        "summary", "description", "status", "comment",
        "priority", "assignee", "reporter",
        "created", "updated", "duedate",
        "attachment",
    ]
    if approver_field:
        fields_list.append(approver_field)
    if reviewer_field:
        fields_list.append(reviewer_field)
    fields_param = ",".join(fields_list)

    # ---- Helper: extract ticket doc from a raw Jira issue ----
    def extract_ticket_doc(issue):
        fields = issue.get("fields", {})
        description = extract_text_from_adf(fields.get("description"))
        comments = [
            extract_text_from_adf(c.get("body"))
            for c in fields.get("comment", {}).get("comments", [])
        ]
        history = []
        for entry in issue.get("changelog", {}).get("histories", []):
            for item in entry.get("items", []):
                if item.get("field") == "status":
                    history.append(f"{item.get('fromString')} -> {item.get('toString')}")

        assignee = fields.get("assignee")
        reporter = fields.get("reporter")
        priority = fields.get("priority")

        approver_val = fields.get(approver_field) if approver_field else None
        approver_name = (
            approver_val.get("displayName") or approver_val.get("name")
            if isinstance(approver_val, dict) else (str(approver_val) if approver_val else None)
        )
        reviewer_val = fields.get(reviewer_field) if reviewer_field else None
        reviewer_name = (
            reviewer_val.get("displayName") or reviewer_val.get("name")
            if isinstance(reviewer_val, dict) else (str(reviewer_val) if reviewer_val else None)
        )
        raw_attachments = fields.get("attachment", [])
        attachments = [
            {
                "attachment_id": att.get("id"),
                "filename": att.get("filename"),
                "mimeType": att.get("mimeType"),
                "content": att.get("content"),  # download URL
            }
            for att in (raw_attachments or [])
            if att.get("id") and att.get("content")
        ]
        return {
            "key": issue.get("key"),
            "summary": fields.get("summary"),
            "description": description,
            "status": fields.get("status", {}).get("name"),
            "priority": priority.get("name") if priority else None,
            "assignee": assignee.get("displayName") if assignee else None,
            "reporter": reporter.get("displayName") if reporter else None,
            "comments": comments,
            "status_history": history,
            "created": fields.get("created"),
            "updated": fields.get("updated"),
            "duedate": fields.get("duedate"),
            "approver": approver_name,
            "reviewer": reviewer_name,
            "attachments": attachments,
        }

    extracted_tickets = []

    # ---- Single-ticket mode ----
    if state.get("ticket_id"):
        try:
            res = requests.get(
                f"{domain}/rest/api/3/issue/{state['ticket_id']}",
                headers={"Accept": "application/json"},
                auth=HTTPBasicAuth(email, token),
                params={"fields": fields_param, "expand": "changelog"},
                timeout=30
            )
            res.raise_for_status()
            extracted_tickets.append(extract_ticket_doc(res.json()))
        except Exception as e:
            logger.error(f"[Fetch] Single ticket fetch failed: {e}")
            return {"final_error": f"Failed to fetch ticket: {e}", "tickets": []}

        return {
            "tickets": extracted_tickets,
            "final_error": "" if extracted_tickets else "Ticket not found."
        }

    # ---- Project mode: resolve board ID dynamically ----
    if not state.get("project_key"):
        return {"final_error": "No Project Key or Ticket ID provided", "tickets": []}

    try:
        board_res = requests.get(
            f"{domain}/rest/agile/1.0/board",
            headers={"Accept": "application/json"},
            auth=HTTPBasicAuth(email, token),
            params={"projectKeyOrId": state["project_key"]},
            timeout=30
        )
        board_res.raise_for_status()
        boards = board_res.json().get("values", [])
        if not boards:
            return {"final_error": f"No board found for project '{state['project_key']}'", "tickets": []}
        board_id = boards[0]["id"]
        logger.info(f"[Fetch] Resolved board ID: {board_id} for project '{state['project_key']}'")
    except Exception as e:
        logger.error(f"[Fetch] Board lookup failed: {e}")
        return {"final_error": f"Failed to resolve board: {e}", "tickets": []}

    # ---- Paginated fetch via Agile Board API ----
    start_at = 0
    max_results = 100

    while True:
        try:
            logger.info(f"[Fetch] Page request — startAt={start_at}, maxResults={max_results}")
            response = requests.get(
                f"{domain}/rest/agile/1.0/board/{board_id}/issue",
                headers={"Accept": "application/json"},
                auth=HTTPBasicAuth(email, token),
                params={
                    "startAt": start_at,
                    "maxResults": max_results,
                    "fields": fields_param,
                    "expand": "changelog"
                },
                timeout=60
            )
            response.raise_for_status()
            data = response.json()

            issues = data.get("issues", [])
            if not issues:
                logger.info("[Fetch] No more issues — pagination complete.")
                break

            # ---- Extract ticket documents using helper ----
            for issue in issues:
                try:
                    extracted_tickets.append(extract_ticket_doc(issue))
                except Exception as e:
                    logger.warning(f"[Fetch] Skipping {issue.get('key', '?')} — extraction error: {e}")
                    continue

            fetched = len(extracted_tickets)
            total_available = data.get("total", 0)
            logger.info(f"[Fetch] Page done — fetched so far: {fetched} / {total_available}")

            start_at += max_results
            if start_at >= total_available:
                logger.info(f"[Fetch] All {total_available} tickets fetched.")
                break

        except Exception as e:
            logger.error(f"[Fetch] Page startAt={start_at} failed: {e}")
            if not extracted_tickets:
                return {"final_error": f"Failed to fetch tickets: {e}", "tickets": []}
            logger.warning(f"[Fetch] Partial recovery — continuing with {len(extracted_tickets)} tickets.")
            break

    logger.info(f"[Fetch] Complete. Total extracted: {len(extracted_tickets)}")
    return {
        "tickets": extracted_tickets,
        "final_error": "" if extracted_tickets else "No tickets found matching criteria."
    }


# -------------------- NODE 2: ANALYZE (ASYNC + SEMAPHORE) --------------------

SYSTEM_PROMPT = """
You are an AI Change Management Compliance Auditor.
Your task is to analyze the provided Jira Ticket JSON and extract specific compliance data.

Focus on identifying:
1.  **Logical Alignment**: Check if the status matches the comments and description history.
2.  **Content Consistency**: Check if the Title (Summary), Description, and Comments are consistent with each other.
3.  **Severity**: Determine severity label (Critical, High, Medium, Low) based on the description content.
4.  **Field Check**: Check if "Due Date", "Approver", and "Reviewer" are set. If missing, list them in `compliance_gaps` but DO NOT automatically mark as "Misaligned" based solely on this. Alignment is determined by logical consistency.
5.  **Compliance Checklist**: Review the provided checklist items. For each item, evaluate if the ticket complies. Provide a 'status' (Pass/Fail/NA) and a 'comment' explaining the evidence.

----------------------------------------
ALIGNMENT RULES
----------------------------------------

Aligned:
- Summary aligns with the description and comments.
- No logical contradictions.
- Workflow progression is consistent.

Misaligned:
- Status is "Done" but work is described as ongoing.
- Status is "Done" but critical governance (testing, approval) is missing.
- Clear contradictions between fields.
- Title, Description, and Comments contradict each other (e.g., Title says "Fix Bug", Description says "New Feature").

Partially Aligned:
- Minor inconsistencies but generally logical.

----------------------------------------
SEVERITY CRITERIA
----------------------------------------

1. **Critical (S1)**: Complete system outage or major security breach. "Core" functionality down for all users.
   - Keywords: Outage, Down, Breach, Data Loss, All users.

2. **High (S2)**: Significant feature failure, no easy workaround. Major part broken, system not dead.
   - Keywords: Broken, Error 500, Cannot export, Critical bug, Not working, Failing.

3. **Medium (S3)**: Partial failure or functional issue with workaround.
   - Keywords: Issue, Workaround, Incorrect display, UI bug.

4. **Low (S4)**: Cosmetic issues or minor inconveniences. No loss of functionality.
   - Keywords: Typo, Suggestion, UI tweak, Enhancement, Color.

----------------------------------------
OUTPUT JSON FORMAT — Return ONLY valid JSON, no markdown, no extra text.
----------------------------------------
{
    "alignment_status": "Aligned | Partially Aligned | Misaligned",
    "severity": "Critical | High | Medium | Low",
    "completion_percentage": 0-100,
    "compliance_gaps": ["List specific missing elements or contradictions"],
    "recommended_actions": ["List specific actions to fix gaps"],
    "compliance_checklist": [
        {
            "category": "Category Name",
            "check": "Specific Check Requirements",
            "status": "Pass | Fail | N/A",
            "comment": "Evidence or reason"
        }
    ]
}
"""

# Max concurrent LLM calls — prevents API rate-limit errors
_SEMAPHORE_LIMIT = 10


async def _analyze_single_ticket(
    ticket: dict,
    llm,
    checklist_str: str,
    semaphore: asyncio.Semaphore,
    tenant_id: str = "anonymous",
    jira_email: str = "",
    jira_token: str = "",
) -> tuple:
    """
    Async analysis of a single ticket, guarded by a semaphore.
    Returns (dashboard_row, is_satisfied, missing_fields). Never raises.
    """
    ticket_key = ticket.get("key", "UNKNOWN")

    # ----------------------------------------------------------------
    # DB cache check — skip LLM if ticket content or attachments changed
    # ----------------------------------------------------------------
    loop = asyncio.get_event_loop()

    # Extract attachment IDs already known from the ticket payload so they
    # can be folded into the ticket hash. A new upload always gets a new
    # attachment_id in Jira, so this automatically invalidates the cache.
    attachment_ids = sorted(
        att.get("attachment_id", "") or ""
        for att in (ticket.get("attachments") or [])
    )
    current_hash = compute_ticket_hash(ticket, attachment_ids)

    existing_hash = await loop.run_in_executor(
        None, get_ticket_hash, tenant_id, ticket_key
    )

    if existing_hash and existing_hash == current_hash:
        cached = await loop.run_in_executor(
            None, load_report_from_db, tenant_id, ticket_key
        )
        if cached:
            logger.info(f"[Analyze] Skipping {ticket_key} — no changes detected (cache hit)")
            is_satisfied = cached.get("is_satisfied", False)
            missing = [
                g for g in cached.get("compliance_gaps", [])
                if not g.startswith("Missing:")
            ] or ([] if is_satisfied else ["(cached gaps — see report)"])
            return cached, is_satisfied, missing

    async with semaphore:
        try:
            # ---- Trim ticket payload to prevent token overflow ----
            minimal_ticket = {
                "key": ticket.get("key"),
                "summary": ticket.get("summary"),
                "description": (ticket.get("description") or "")[:1500],
                "status": ticket.get("status"),
                "priority": ticket.get("priority"),
                "assignee": ticket.get("assignee"),
                "approver": ticket.get("approver"),
                "reviewer": ticket.get("reviewer"),
                "duedate": ticket.get("duedate"),
                "status_history": ticket.get("status_history", [])[:5],
                "comments": [c[:300] for c in (ticket.get("comments") or [])[:3]],
            }
            ticket_str = json.dumps(minimal_ticket)
            user_prompt = (
                f"Analyze this ticket:\n{ticket_str}\n\n"
                f"Verify against the following checklist:\n{checklist_str}"
            )

            response = await llm.ainvoke([
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt}
            ])

            content = response.content if hasattr(response, "content") else str(response)
            analysis = safe_json_parse(content)

            if not analysis:
                logger.warning(f"[Analyze] {ticket_key} — JSON parse failed. Raw: {content[:200]}")
                analysis = _error_analysis("Could not extract JSON from LLM response")

        except Exception as e:
            logger.error(f"[Analyze] {ticket_key} — LLM call failed: {e}")
            analysis = _error_analysis(str(e))

    # ---- Strict compliance field checks (always runs, even on error) ----
    missing_fields = []

    # 1. Documentation completeness
    if not ticket.get("summary"):       missing_fields.append("Summary")
    if not ticket.get("description"):   missing_fields.append("Description")
    if not ticket.get("priority"):      missing_fields.append("Priority")

    # 2. Audit trail
    if not ticket.get("assignee"):      missing_fields.append("Assignee")
    if not ticket.get("created"):       missing_fields.append("Created_Date")
    if not ticket.get("duedate"):       missing_fields.append("Due_Date")

    # 3. Governance — Approver and Reviewer are mandatory
    if not ticket.get("approver"):      missing_fields.append("Approver")
    if not ticket.get("reviewer"):      missing_fields.append("Reviewer")

    # 4. LLM logical alignment
    llm_alignment = analysis.get("alignment_status", "Misaligned")
    if llm_alignment not in ("Aligned",):
        missing_fields.append(f"Logical Misalignment ({llm_alignment})")

    # 5. Compliance checklist — any "Fail" item makes ticket dissatisfied
    checklist_results = analysis.get("compliance_checklist", [])
    failed_checks = [
        f"Checklist Fail: [{item.get('category', '')}] {item.get('check', '')}"
        for item in checklist_results
        if str(item.get("status", "")).strip().lower() == "fail"
    ]
    if failed_checks:
        missing_fields.extend(failed_checks)

    is_satisfied = len(missing_fields) == 0

    # ---- Build compliance_gaps for dashboard ----
    if not is_satisfied:
        llm_gaps = analysis.get("compliance_gaps", [])
        hard_missing = [f for f in missing_fields if "Logical" not in f and "Checklist Fail" not in f]
        checklist_fails = [f for f in missing_fields if "Checklist Fail" in f]

        parts = []
        if hard_missing:
            parts.append(f"Missing: {', '.join(hard_missing)}")
        parts.extend(checklist_fails)
        parts.extend(llm_gaps)
        final_gaps = parts
    else:
        final_gaps = []

    # ----------------------------------------------------------------
    # 6. Attachment relevance check (embedding-based, cached by attachment_id)
    # ----------------------------------------------------------------
    attachments = ticket.get("attachments", []) or []
    ticket_priority = ticket.get("priority", "") or ""
    att_summary: dict = {"total": 0, "relevant": 0, "irrelevant": 0, "scores": {}}

    if jira_email and jira_token and attachments:
        ticket_context = " ".join(filter(None, [
            ticket.get("summary", ""),
            ticket.get("description", ""),
            " ".join(ticket.get("comments", []) or []),
        ]))

        try:
            att_result = await loop.run_in_executor(
                None,
                process_ticket_attachments,
                attachments,
                jira_email,
                jira_token,
                ticket_context,
                tenant_id,
                ticket_key,
            )
            irrelevant_details = [
                {"filename": d["filename"], "score": d["score"]}
                for d in att_result.get("details", [])
                if not d["is_relevant"]
            ]
            att_summary = {
                "total": att_result["total"],
                "relevant": att_result["relevant"],
                "irrelevant": att_result["irrelevant"],
                "irrelevant_files": irrelevant_details,
                "scores": {
                    d["filename"]: d["score"]
                    for d in att_result.get("details", [])
                },
            }

            # If ALL attachments are irrelevant → compliance gap
            if att_result["total"] > 0 and not att_result["any_relevant"]:
                irrelevant_names = ", ".join(
                    d["filename"] for d in att_result.get("details", [])
                    if not d["is_relevant"]
                )
                final_gaps.append(
                    f"Attachment not relevant to ticket content: {irrelevant_names}"
                )
                is_satisfied = False
                logger.info(f"[Analyze] {ticket_key} — all attachments irrelevant, marked Dissatisfied")

        except Exception as att_exc:
            logger.warning(f"[Analyze] {ticket_key} — attachment check failed (non-fatal): {att_exc}")

    # Edge case: High/Critical with NO attachments → mandatory attachment gap
    if ticket_priority.lower() in ("high", "critical") and not attachments:
        final_gaps.append("Mandatory attachment missing for high-severity ticket")
        is_satisfied = False
        logger.info(f"[Analyze] {ticket_key} — High/Critical priority with no attachments")

    dashboard_row = {
        "key": ticket.get("key"),
        "summary": ticket.get("summary"),
        "status": ticket.get("status"),
        "assignee": ticket.get("assignee"),
        "received_priority": ticket.get("priority"),
        "priority": analysis.get("severity"),
        "alignment_status": llm_alignment,
        "completion_percentage": analysis.get("completion_percentage", 0),
        "compliance_gaps": final_gaps,
        "recommended_actions": analysis.get("recommended_actions", []),
        "compliance_checklist": checklist_results,
        "is_satisfied": is_satisfied,
        "attachment_check": att_summary,
    }

    # ----------------------------------------------------------------
    # Persist to DB (upsert — safe to call multiple times)
    # ----------------------------------------------------------------
    try:
        await loop.run_in_executor(
            None, upsert_ticket,
            tenant_id, ticket_key,
            "UNKNOWN",
            ticket.get("updated"),
            current_hash,
        )
        await loop.run_in_executor(
            None, upsert_report,
            tenant_id, ticket_key, dashboard_row,
        )
    except Exception as db_exc:
        logger.warning(f"[DB] Upsert failed for {ticket_key}: {db_exc} (analysis still returned)")

    return dashboard_row, is_satisfied, missing_fields


def analyze_compliance(state: ComplianceState):
    if state.get("final_error"):
        return {"project_report": []}

    tickets = state["tickets"]
    total = len(tickets)
    logger.info(f"[Analyze] Starting analysis of {total} tickets.")

    llm = get_llm(
        temperature=0
    )

    # Read compliance checklist once
    checklist_data = read_compliance_table()
    checklist_str = ""
    if checklist_data:
        checklist_str = "COMPLIANCE CHECKLIST:\n"
        for i, item in enumerate(checklist_data, 1):
            checklist_str += f"{i}. Category: {item['category']} - Check: {item['check']}\n"

    tenant_id: str = state.get("tenant_id", "anonymous")
    jira_email: str = state.get("jira_email", "")
    jira_token: str = state.get("jira_token", "")

    # ---- Async runner with semaphore ----
    async def run_all_async():
        semaphore = asyncio.Semaphore(_SEMAPHORE_LIMIT)
        tasks = [
            _analyze_single_ticket(
                ticket, llm, checklist_str, semaphore,
                tenant_id, jira_email, jira_token,
            )
            for ticket in tickets
        ]
        return await asyncio.gather(*tasks, return_exceptions=True)

    # ---- Execute — handle whether we're inside an existing event loop or not ----
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # LangGraph worker thread: run in a new thread with its own loop
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                future = pool.submit(asyncio.run, run_all_async())
                raw_results = future.result(timeout=3600)
        else:
            raw_results = loop.run_until_complete(run_all_async())
    except RuntimeError:
        # No event loop at all — create one
        raw_results = asyncio.run(run_all_async())

    # ---- Process results ----
    dashboard_results = []
    satisfied_tickets = []
    dissatisfied_tickets = []

    for ticket, raw in zip(tickets, raw_results):
        ticket_key = ticket.get("key", "?")

        if isinstance(raw, Exception):
            logger.error(f"[Analyze] {ticket_key} — unhandled exception: {raw}")
            dashboard_results.append({
                "key": ticket_key,
                "summary": ticket.get("summary"),
                "status": "Error",
                "alignment_status": "Error",
                "completion_percentage": 0,
                "compliance_gaps": [str(raw)],
                "recommended_actions": ["Retry analysis"],
                "compliance_checklist": [],
                "is_satisfied": False,
            })
            continue

        dashboard_row, is_satisfied, missing_fields = raw
        dashboard_results.append(dashboard_row)

        if is_satisfied:
            satisfied_tickets.append({
                "id": ticket_key,
                "summary": ticket.get("summary"),
                "status": "Passed"
            })
        else:
            dissatisfied_tickets.append({
                "id": ticket_key,
                "status": "❌ Dissatisfied",
                "missing_fields": missing_fields,
                "deficiency": (
                    f"Ticket failed validation. "
                    f"Issues: {', '.join(missing_fields[:5])}. "
                    + (dashboard_row["compliance_gaps"][0] if dashboard_row["compliance_gaps"] else "")
                )
            })

    # ---- Generate Markdown Report ----
    md_report = "# Ticket Compliance Audit Report\n\n"

    md_report += "## 1. Satisfied\n"
    if satisfied_tickets:
        md_report += "| Ticket ID | Summary | Status |\n|---|---|---|\n"
        for t in satisfied_tickets:
            md_report += f"| {t['id']} | {t['summary']} | {t['status']} |\n"
    else:
        md_report += "*No tickets satisfied all criteria.*\n"

    md_report += "\n## 2. Dissatisfied\n"
    if dissatisfied_tickets:
        for t in dissatisfied_tickets:
            md_report += f"### Ticket ID: {t['id']}\n"
            md_report += f"- **Status**: {t['status']}\n"
            md_report += f"- **Missing/Incomplete Fields**: {', '.join(t['missing_fields'])}\n"
            md_report += f"- **Detailed Deficiency**: {t['deficiency']}\n\n"
    else:
        md_report += "*No dissatisfied tickets found.*\n"

    logger.info(
        f"[Analyze] Complete — Total={total}, "
        f"Satisfied={len(satisfied_tickets)}, "
        f"Dissatisfied={len(dissatisfied_tickets)}"
    )

    return {"project_report": dashboard_results}


# -------------------- GRAPH --------------------

workflow = StateGraph(ComplianceState)

workflow.add_node("fetch", fetch_tickets)
workflow.add_node("analyze", analyze_compliance)

workflow.set_entry_point("fetch")
workflow.add_edge("fetch", "analyze")
workflow.add_edge("analyze", END)

app = workflow.compile()


# -------------------- ENTRY FUNCTION --------------------

async def run_compliance_agent(
    jira_domain: str,
    jira_email: str,
    jira_token: str,
    project_key: str = None,
    ticket_id: str = None,
    tenant_id: str = "anonymous",
):
    final_state = await app.ainvoke({
        "project_key": project_key,
        "ticket_id": ticket_id,
        "jira_domain": jira_domain,
        "jira_email": jira_email,
        "jira_token": jira_token,
        "tickets": [],
        "project_report": [],
        "final_error": "",
        "tenant_id": tenant_id,
    })

    report = final_state.get("project_report", [])
    error = final_state.get("final_error")

    if error and not report:
        return {
            "status": "Analysis Failed",
            "analysis_result": {
                "status": "Error",
                "details": error
            }
        }

    return {
        "status": "Project Analysis Completed",
        "analysis_result": report
    }
